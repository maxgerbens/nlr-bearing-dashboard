"""
Migreert v0.2 naar v0.3.

Wijziging: paragrafen 4.1 (Functionele eisen), 4.2 (Niet-functionele eisen) en
4.3 (User stories) worden vervangen door een gekoppelde versie waarin eisen
ID's krijgen (FE-1..7, NFE-1..3) en elke user story expliciet naar die ID's
verwijst, plus een dekking-matrix onderaan.

De rest van het document (eigen aanvullingen in 1.1, 2.x, etc.) blijft
ongewijzigd. Versielabel op de voorpagina wordt bijgewerkt naar 0.3.

Let op: na openen in Word de inhoudsopgave verversen (rechtermuisknop ->
'Veld bijwerken') zodat de paginanummers kloppen.
"""

from docx import Document

SRC = "/Users/maxgerbens/Documents/GitHub/nlr-bearing-dashboard/docs/Ontwerpdocument_NLR_Bearing_Dashboard_v0.2.docx"
DST = "/Users/maxgerbens/Documents/GitHub/nlr-bearing-dashboard/docs/Ontwerpdocument_NLR_Bearing_Dashboard_v0.3.docx"


def main():
    doc = Document(SRC)

    # --- 1. Versie bijwerken op voorpagina --------------------------------
    for p in doc.paragraphs:
        for run in p.runs:
            if "Versie: 0.2" in run.text:
                run.text = run.text.replace("Versie: 0.2 (skeleton)", "Versie: 0.3")
                run.text = run.text.replace("Versie: 0.2", "Versie: 0.3")

    # --- 2. Sectie 4.1 t/m 4.3 lokaliseren --------------------------------
    start_idx = None
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "4.1 Functionele eisen":
            start_idx = i
        elif t == "4.4 Architectuur":
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        raise SystemExit("Kon sectie 4.1 of 4.4 niet vinden in v0.2 — script gestopt.")

    anchor_el = doc.paragraphs[end_idx]._element

    # --- 3. Oude 4.1-4.3 paragrafen verwijderen ---------------------------
    for p in list(doc.paragraphs[start_idx:end_idx]):
        p._element.getparent().remove(p._element)

    # --- 4. Helpers voor invoegen vóór de 4.4-anchor ----------------------
    def add_heading_before(text, level):
        p = doc.add_heading(text, level=level)
        anchor_el.addprevious(p._element)

    def add_para_before(runs):
        """runs = list of (text, bold) tuples"""
        p = doc.add_paragraph()
        for text, bold in runs:
            r = p.add_run(text)
            r.bold = bold
        anchor_el.addprevious(p._element)

    def add_table_before(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                table.rows[ri].cells[ci].text = val
        anchor_el.addprevious(table._element)
        spacer = doc.add_paragraph()
        anchor_el.addprevious(spacer._element)

    # --- 5. 4.1 Functionele eisen -----------------------------------------
    add_heading_before("4.1 Functionele eisen", level=2)
    add_para_before([(
        "Elke functionele eis krijgt een ID (FE-1 t/m FE-7), zodat de user "
        "stories in 4.3 expliciet kunnen aangeven welke eis ze invullen.", False
    )])
    add_table_before(
        ["ID", "Functionele eis"],
        [
            ("FE-1", "Het dashboard maakt een lager en belastingsconditie selecteerbaar, zodat de monteur gericht een specifieke meting kan bekijken."),
            ("FE-2", "Het dashboard toont het trendverloop van RMS, kurtosis en crest factor over de levensduur van het geselecteerde lager."),
            ("FE-3", "Het dashboard markeert het moment van falen in de trendgrafiek."),
            ("FE-4", "Het dashboard toont het ruwe vibratiesignaal van een geselecteerd meetmoment."),
            ("FE-5", "Het dashboard maakt lagers onder dezelfde belastingsconditie onderling vergelijkbaar in één grafiek."),
            ("FE-6", "Het dashboard toont het faulttype per lager."),
            ("FE-7", "Het dashboard filtert lagers op belastingsconditie."),
        ],
    )

    # --- 6. 4.2 Niet-functionele eisen ------------------------------------
    add_heading_before("4.2 Niet-functionele eisen", level=2)
    add_para_before([(
        "De niet-functionele eisen krijgen ID's NFE-1 t/m NFE-3 en worden in "
        "de user stories in 4.3 als acceptatievoorwaarde meegenomen.", False
    )])
    add_table_before(
        ["ID", "Niet-functionele eis"],
        [
            ("NFE-1", "Een trendgrafiek verschijnt binnen 2 seconden na selectie van een lager, zodat de monteur niet onnodig wacht."),
            ("NFE-2", "Het dashboard is bedienbaar zonder instructie; alle knoppen en filters hebben duidelijke labels."),
            ("NFE-3", "Elke visualisatie bevat een korte toelichting die uitlegt wat het kenmerk betekent en waarom het relevant is voor slijtage."),
        ],
    )

    # --- 7. 4.3 User stories ----------------------------------------------
    add_heading_before("4.3 User stories", level=2)
    add_para_before([(
        "De stories zijn geordend volgens de natuurlijke werkstroom van de "
        "monteur: eerst de relevante lagers afbakenen (US-1), dan vergelijken "
        "(US-2), vervolgens inzoomen op één lager (US-3 en US-4) en ten slotte "
        "een specifiek meetmoment inspecteren (US-5). Elke story benoemt "
        "expliciet welke eisen uit 4.1 en 4.2 eraan invulling geven; de "
        "dekking-matrix onderaan deze paragraaf maakt zichtbaar dat alle eisen "
        "minstens één keer voorkomen.", False
    )])

    stories = [
        {
            "id": "US-1",
            "titel": "Filteren op belastingsconditie",
            "story": (
                "Als onderhoudsmonteur wil ik kunnen filteren op "
                "belastingsconditie, zodat ik alleen lagers zie die "
                "vergelijkbaar zijn met mijn situatie."
            ),
            "fe": "FE-7, FE-1",
            "nfe": "NFE-1, NFE-2",
            "ac": (
                "Het dashboard bevat een filter met de drie beschikbare "
                "condities (35 Hz/12 kN, 37,5 Hz/11 kN, 40 Hz/10 kN); na "
                "selectie past de lagerlijst en grafiek zich binnen 2 seconden "
                "aan."
            ),
        },
        {
            "id": "US-2",
            "titel": "Lagers onderling vergelijken",
            "story": (
                "Als onderhoudsmonteur wil ik twee of meer lagers onder "
                "dezelfde belastingsconditie naast elkaar zien, zodat ik "
                "afwijkend gedrag kan herkennen."
            ),
            "fe": "FE-5, FE-6, FE-7",
            "nfe": "NFE-2, NFE-3",
            "ac": (
                "Het dashboard toont minimaal twee lagers tegelijkertijd onder "
                "dezelfde belastingsconditie, met het faulttype per lager "
                "zichtbaar; de bediening verloopt via duidelijk gelabelde "
                "filters en selecties zonder externe instructie."
            ),
        },
        {
            "id": "US-3",
            "titel": "Trendverloop van vibratiekenmerken bekijken",
            "story": (
                "Als onderhoudsmonteur wil ik het trendverloop van RMS, "
                "kurtosis en crest factor per lager zien, zodat ik kan "
                "beoordelen of een lager tekenen van degradatie vertoont."
            ),
            "fe": "FE-1, FE-2, FE-6",
            "nfe": "NFE-1, NFE-3",
            "ac": (
                "Na selectie van een lager toont het dashboard binnen 2 "
                "seconden een grafiek met RMS, kurtosis en crest factor over "
                "de volledige levensduur, met het faulttype zichtbaar en een "
                "korte toelichting per kenmerk."
            ),
        },
        {
            "id": "US-4",
            "titel": "Moment van falen zichtbaar maken",
            "story": (
                "Als onderhoudsmonteur wil ik het moment van falen zien in de "
                "trendgrafiek, zodat ik kan beoordelen hoe vroeg de kenmerken "
                "begonnen af te wijken."
            ),
            "fe": "FE-2, FE-3",
            "nfe": "NFE-3",
            "ac": (
                "Het falenmoment is als verticale markering zichtbaar in de "
                "trendgrafiek van elk lager waarvoor faaldata beschikbaar is, "
                "met een toelichting bij de markering."
            ),
        },
        {
            "id": "US-5",
            "titel": "Ruwe meting van een meetmoment inspecteren",
            "story": (
                "Als onderhoudsmonteur wil ik het ruwe vibratiesignaal van "
                "een specifiek meetmoment openen, zodat ik een afwijking op "
                "detailniveau kan controleren."
            ),
            "fe": "FE-1, FE-4",
            "nfe": "NFE-1, NFE-3",
            "ac": (
                "Bij selectie van een punt in de trendgrafiek wordt binnen 2 "
                "seconden het bijbehorende ruwe vibratiesignaal getoond, met "
                "toelichting over de meeteenheid en bemonsteringsfrequentie."
            ),
        },
    ]

    # User stories als één samenhangende tabel
    headers = ["ID", "User story", "Functionele eisen", "Niet-functionele eisen", "Acceptatiecriterium"]
    us_table = doc.add_table(rows=1 + len(stories), cols=len(headers))
    us_table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = us_table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for ri, s in enumerate(stories, start=1):
        row = us_table.rows[ri].cells
        row[0].text = s["id"]
        # Kolom 'User story': bold titel + story-tekst eronder
        story_cell = row[1]
        title_para = story_cell.paragraphs[0]
        title_run = title_para.add_run(s["titel"])
        title_run.bold = True
        story_cell.add_paragraph(s["story"])
        row[2].text = s["fe"]
        row[3].text = s["nfe"]
        row[4].text = s["ac"]
    anchor_el.addprevious(us_table._element)
    spacer = doc.add_paragraph()
    anchor_el.addprevious(spacer._element)

    # --- 8. Dekking-matrix ------------------------------------------------
    add_para_before([("Dekking-matrix (eisen → user stories)", True)])
    add_table_before(
        ["Eis", "Gedekt door"],
        [
            ("FE-1", "US-1, US-3, US-5"),
            ("FE-2", "US-3, US-4"),
            ("FE-3", "US-4"),
            ("FE-4", "US-5"),
            ("FE-5", "US-2"),
            ("FE-6", "US-2, US-3"),
            ("FE-7", "US-1, US-2"),
            ("NFE-1", "US-1, US-3, US-5"),
            ("NFE-2", "US-1, US-2"),
            ("NFE-3", "US-2, US-3, US-4, US-5"),
        ],
    )

    doc.save(DST)
    print(f"Opgeslagen: {DST}")


if __name__ == "__main__":
    main()
