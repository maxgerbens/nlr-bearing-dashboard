"""
Bouwt v0.2 van het ontwerpdocument als .docx.

Wijzigingen t.o.v. v0.1:
- Hoofdstuk 3 (Data Understanding) ingekort tot één paragraaf 'Bron en omvang'.
- Hoofdstuk 5 (Aanpak en planning) verwijderd — niet beoordeeld bij data analytics.
- MC1-rubric-blok verwijderd (MC1 wordt niet beoordeeld bij data analytics).
- Eenvoudige opmaak: standaard Word-stijlen, geen kleuren, geen cursief.
- Renumbering van de overige hoofdstukken.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_prompt(doc, text):
    """Invul-prompt: één alinea, standaard tekst, bold label."""
    p = doc.add_paragraph()
    r = p.add_run("Invullen: ")
    r.bold = True
    p.add_run(text)


def add_rubric(doc, leeruitkomst, op_niveau, boven_niveau):
    """Rubric-hint: één alinea met drie regels, bold labels."""
    p = doc.add_paragraph()
    r = p.add_run(f"Rubric ({leeruitkomst})")
    r.bold = True
    p.add_run("\n")
    r = p.add_run("Op niveau: ")
    r.bold = True
    p.add_run(op_niveau + "\n")
    r = p.add_run("Boven niveau: ")
    r.bold = True
    p.add_run(boven_niveau)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Rechtermuisknop -> 'Veld bijwerken' om inhoudsopgave te genereren."
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def build():
    doc = Document()
    setup_styles(doc)

    # ----- Voorblad -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Ontwerpdocument")
    r.bold = True
    r.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("NLR Bearing Dashboard")
    r.bold = True
    r.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.add_run("Differentiatieopdracht DP10 - Data Analytics")

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Auteur: [naam student]\n")
    meta.add_run("Studentnummer: [studentnummer]\n")
    meta.add_run("Opleiding: HBO-ICT / Data Science, Hogeschool Windesheim\n")
    meta.add_run("Docent: [naam docent]\n")
    meta.add_run("Datum: [datum oplevering]\n")
    meta.add_run("Versie: 0.2 (skeleton)")

    doc.add_page_break()

    # ----- Inhoudsopgave -----
    doc.add_heading("Inhoudsopgave", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ----- Samenvatting -----
    doc.add_heading("Managementsamenvatting", level=1)
    add_prompt(doc, "Een halve tot één pagina. Schrijf deze sectie als laatste. "
                    "Vat aanleiding, businessvraag, gekozen oplossing (dashboard), "
                    "belangrijkste keuzes en verwachte waarde voor NLR samen.")
    doc.add_page_break()

    # ===== 1. Inleiding =====
    doc.add_heading("1. Inleiding", level=1)

    doc.add_heading("1.1 Aanleiding", level=2)
    add_prompt(doc,
        "Beschrijf de fictieve casus: NLR (Nederlands Lucht- en Ruimtevaartcentrum) "
        "wil de bedrijfszekerheid van roterende componenten vergroten via predictive "
        "maintenance op lagers. Leg uit waarom dit relevant is (veiligheid, kosten van "
        "ongeplande stilstand, duurzaamheid). Vermeld dat NLR hier als fictieve "
        "opdrachtgever fungeert en dat de publieke XJTU-SY dataset als databron dient.")

    doc.add_heading("1.2 Doel van dit document", level=2)
    add_prompt(doc,
        "Eén alinea. Dit document beschrijft de business- en ontwerpkeuzes voor een "
        "interactief dashboard dat lagergegevens analyseert en presenteert aan een "
        "specifieke eindgebruiker bij NLR.")

    doc.add_page_break()

    # ===== 2. Business Understanding (AN1) =====
    doc.add_heading("2. Business Understanding", level=1)
    add_rubric(doc, "AN1 Business Understanding",
        "Heldere businessdoelstelling en -vraag, organisatorische context, stakeholders, "
        "juridische implicaties (AVG, AI Act), concrete data-mining doelen.",
        "Scherpe, kritisch geformuleerde businessvraag; afstemming met stakeholders; "
        "diepgaande reflectie op organisatorische en maatschappelijke context; "
        "juridische implicaties leiden aantoonbaar tot bewuste keuzes; "
        "meetbare data-miningdoelen vertaald naar KPI's en kritische succesfactoren.")

    doc.add_heading("2.1 Organisatorische context", level=2)
    add_prompt(doc,
        "Beschrijf NLR kort (rol in NL luchtvaartsector, R&D-karakter). Plaats predictive "
        "maintenance binnen NLR: bijvoorbeeld testopstellingen, MRO-onderzoek of advies "
        "aan luchtvaartmaatschappijen. Verwijs naar publieke info over NLR; geen verzonnen "
        "interne details.")

    doc.add_heading("2.2 Stakeholders en eindgebruiker", level=2)
    doc.add_paragraph(
        "Het dashboard moet één primaire eindgebruiker dienen. Hieronder drie kandidaat-"
        "rollen met voor- en nadelen. Kies er één en motiveer.")

    headers = ["Eindgebruiker", "Beslissing die dashboard ondersteunt", "Voordelen", "Nadelen"]
    rows = [
        ("Onderhoudsmonteur / maintenance engineer",
         "Wanneer moet dit lager preventief vervangen worden?",
         "Direct operationeel; concrete waarde; eenvoudige KPI's.",
         "Eist hoge nauwkeurigheid en uitlegbaarheid; veiligheidsimplicaties bij fout."),
        ("Reliability / data engineer",
         "Welke faulttypes komen voor en hoe ontwikkelen features zich richting falen?",
         "Past natuurlijk bij dataset; ruimte voor vergelijking van features en condities; "
         "minder veiligheidskritisch.",
         "Doelgroep is technisch; businesswaarde moet expliciet onderbouwd worden."),
        ("R&D-onderzoeker bij NLR",
         "Welke signaalkenmerken zijn meest indicatief voor degradatie onder verschillende "
         "operating conditions?",
         "Past bij R&D-karakter NLR; legitimeert verkennend dashboard; sluit aan bij "
         "academische dataset.",
         "Businesswaarde minder direct; lastiger meetbare KPI's."),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val

    add_prompt(doc,
        "Kies één primaire eindgebruiker en motiveer in één of twee alinea's waarom. "
        "Benoem ook secundaire stakeholders (bv. teamlead, IT-beheer NLR). Voor 'boven "
        "niveau': beschrijf hoe je deze keuze zou afstemmen met de stakeholder, ook al "
        "is dat in de fictieve casus niet fysiek mogelijk.")

    doc.add_heading("2.3 Maatschappelijke ontwikkelingen", level=2)
    add_prompt(doc,
        "Plaats de opdracht in bredere context: predictive maintenance reduceert "
        "ongeplande stilstand en materiaalverspilling; luchtvaartveiligheid; groeiende "
        "rol van AI in safety-kritische omgevingen. Verwijs vooruit naar hoofdstuk 6.")

    doc.add_heading("2.4 Businessdoelstelling", level=2)
    add_prompt(doc,
        "Eén heldere zin volgens SMART. Voorbeeld: 'NLR wil binnen [tijd] de "
        "[onderhoudsbeslissing] van [doelgroep] ondersteunen met een dashboard dat "
        "[meetbaar effect] realiseert.' Vermijd vage termen als 'inzicht geven'.")

    doc.add_heading("2.5 Businessvraag", level=2)
    add_prompt(doc,
        "Formuleer de hoofdvraag zo dat het antwoord een beslissing ondersteunt. "
        "Niet: 'Wat zijn kenmerken van falende lagers?' (beschrijvend). "
        "Wel: 'Welke combinatie van vibratiekenmerken voorspelt het einde van de "
        "bruikbare levensduur van een UER204-lager onder operating condition X met "
        "voldoende vooruitzicht voor planning van onderhoud?'")

    doc.add_heading("2.6 Juridische implicaties (AVG en AI Act)", level=2)
    add_prompt(doc,
        "Behandel expliciet: (a) AVG - bevat de dataset persoonsgegevens? Nee, het zijn "
        "machinemetingen; leg uit waarom dat de conclusie is en wat er zou veranderen "
        "als NLR het in productie zou nemen met operator-IDs. (b) EU AI Act - "
        "classificeer het systeem. Een onderhoudsbeslissing in luchtvaart kan onder "
        "'high-risk' vallen; bespreek welke verplichtingen dat geeft (transparantie, "
        "menselijk toezicht, robuustheid, documentatie). Voor 'boven niveau': benoem "
        "welke ontwerpkeuzes in hoofdstuk 4 hieruit voortvloeien.")

    doc.add_heading("2.7 Data-mining doelen, KPI's en succesfactoren", level=2)
    add_prompt(doc,
        "Vertaal de businessvraag naar twee tot vier concrete, meetbare data-mining "
        "doelen. Voorbeeld: 'Classificeer faulttype (4 klassen) met minimaal X% accuracy'; "
        "'Schat Remaining Useful Life met MAE kleiner dan Y% van werkelijke levensduur'. "
        "Koppel elk doel aan een KPI op het dashboard en een kritische succesfactor "
        "(bv. gebruiker begrijpt visualisatie binnen 30 sec). Voor 'boven niveau' is "
        "meetbaarheid plus koppeling aan businessdoel verplicht.")

    doc.add_page_break()

    # ===== 3. Data (kort) =====
    doc.add_heading("3. Databron", level=1)
    doc.add_paragraph(
        "Dit project gebruikt de publieke XJTU-SY Bearing Dataset (Wang et al., 2018, "
        "DOI: 10.1109/TR.2018.2882682), samengesteld door Xi'an Jiaotong University en "
        "Changxing Sumyoung Technology. De dataset bevat run-to-failure-vibratiedata van "
        "15 rolling element bearings (type LDK UER204), gemeten op twee assen met 25,6 "
        "kHz. De lagers zijn verdeeld over drie operating conditions (35 Hz/12 kN, "
        "37,5 Hz/11 kN, 40 Hz/10 kN), vijf lagers per conditie. De gedocumenteerde "
        "faulttypes zijn outer race, inner race, cage en ball - of combinaties daarvan."
    )
    add_prompt(doc,
        "Voeg eventueel toe: welke beperkingen van de dataset relevant zijn voor de "
        "businessvraag en data-mining doelen uit hoofdstuk 2 (bv. klein aantal lagers "
        "per conditie, enkel UER204, geen procesparameters zoals temperatuur). Koppel "
        "elke beperking aan een keuze of voorbehoud in dit document.")

    doc.add_page_break()

    # ===== 4. Ontwerp =====
    doc.add_heading("4. Ontwerp van het dashboard", level=1)

    doc.add_heading("4.1 Functionele eisen", level=2)
    add_prompt(doc,
        "Lijst van wat het dashboard moet kunnen, afgeleid van de gekozen eindgebruiker "
        "(par. 2.2). Schrijf als 'Het dashboard moet [actie] zodat [eindgebruiker] "
        "[beslissing]'. Minimaal vijf tot acht eisen.")

    doc.add_heading("4.2 Niet-functionele eisen", level=2)
    add_prompt(doc,
        "Performance (bv. laadtijd grafiek korter dan 2 sec), gebruiksvriendelijkheid, "
        "onderhoudbaarheid, security/privacy (zie par. 2.6), uitlegbaarheid van "
        "modeluitvoer.")

    doc.add_heading("4.3 User stories", level=2)
    add_prompt(doc,
        "Twee tot vier user stories: 'Als [rol] wil ik [actie] zodat [waarde]'. Voeg "
        "per story een acceptatiecriterium toe.")

    doc.add_heading("4.4 Architectuur", level=2)
    add_prompt(doc,
        "Schets de datastroom: ruwe CSV (data/) -> preprocessing -> feature extractie "
        "(features/) -> aggregatie -> visualisatielaag -> dashboard-UI. Benoem per stap: "
        "wat gebeurt er, welk artefact wordt opgeslagen, welke tools.")

    doc.add_heading("4.5 Technologiekeuze", level=2)
    add_prompt(doc,
        "Maak en motiveer keuzes voor: programmeertaal, dashboard-tool (Streamlit / "
        "Dash / Power BI / anders), feature-bibliotheek, opslagformaat. Vergelijk "
        "minimaal twee alternatieven per keuze.")

    doc.add_heading("4.6 Schermontwerp", level=2)
    add_prompt(doc,
        "Voeg minimaal één wireframe of mock-up toe. Beschrijf per scherm: welke vraag "
        "het beantwoordt, welke visualisaties erop staan, welke interactie mogelijk is.")

    doc.add_heading("4.7 KPI's en visualisaties", level=2)
    add_prompt(doc,
        "Per data-miningdoel uit par. 2.7: welke visualisatie toont het, welke KPI "
        "ernaast.")

    doc.add_page_break()

    # ===== 5. Deployment (was 6) =====
    doc.add_heading("5. Deployment en overdracht", level=1)
    add_rubric(doc, "AD2 Deployment",
        "Final report presenteert resultaten, vervolgstappen en proces; data product is overgedragen.",
        "Detailniveau passend bij opdrachtgever; ook kosten, afwijkingen, implementatieplan "
        "en aanbevelingen voor toekomstige projecten.")

    doc.add_heading("5.1 Wijze van oplevering", level=2)
    add_prompt(doc,
        "Hoe wordt het dashboard opgeleverd? Bv. lokale Streamlit-app plus repository "
        "plus README; geen productie-hosting binnen scope. Wees realistisch over de "
        "fictieve casus.")

    doc.add_heading("5.2 Overdrachtsmoment en -materiaal", level=2)
    add_prompt(doc,
        "Wat wordt bij overdracht aan NLR meegeleverd: dit ontwerpdocument, broncode-"
        "repository, README, handleiding, final report.")

    doc.add_heading("5.3 Kosten en afwijkingen", level=2)
    add_prompt(doc,
        "Voor 'boven niveau' vraagt de rubric expliciet aandacht voor gemaakte kosten "
        "(urenraming) en afwijkingen t.o.v. oorspronkelijk plan. Schrijf in deze "
        "fictieve casus over geschatte studie-uren en welke geplande features wel/niet "
        "haalbaar bleken.")

    doc.add_heading("5.4 Aanbevelingen voor vervolg", level=2)
    add_prompt(doc,
        "Concreet en onderbouwd: welke data zou NLR nog moeten verzamelen "
        "(procesparameters, meer lagertypes)? Welke vervolgmodellen? Wanneer is "
        "doorzetten naar productie zinvol of juist niet?")

    doc.add_page_break()

    # ===== 6. Ethiek (was 7) =====
    doc.add_heading("6. Ethische en maatschappelijke overwegingen", level=1)
    add_rubric(doc, "AD3 Ethiek & Maatschappij (beoordeling op de bijlage)",
        "Maatschappelijke implicaties benoemd; relevante morele vraag; moreel oordeel "
        "onderbouwd vanuit verschillende ethische invalshoeken; kritische reflectie.",
        "Maatschappelijke implicaties leiden aantoonbaar tot bewuste keuzes; dilemma "
        "duidelijk in kaart; oordeel gebaseerd op heldere morele uitgangspunten; "
        "reflectie aan de hand van beantwoorde gewetensvragen.")

    add_prompt(doc,
        "Dit hoofdstuk geeft de samenvatting; het volledige moreel stappenplan staat in "
        "de aparte bijlage 'Ethische toelichting'. Vat hier in één pagina samen: "
        "maatschappelijke implicaties (veiligheid, duurzaamheid, vertrouwen in AI); "
        "de morele kernvraag (bv. 'Mag een dashboardadvies over lagervervanging leidend "
        "zijn boven het oordeel van de monteur?'); welke ontwerpkeuzes in hoofdstuk 4 "
        "hier direct uit voortvloeien (uitlegbaarheid, human-in-the-loop, "
        "onzekerheidsweergave).")

    doc.add_page_break()

    # ===== Bijlagen =====
    doc.add_heading("Bijlagen", level=1)
    doc.add_paragraph("A. Ethische toelichting + moreel stappenplan (apart document)")
    doc.add_paragraph("B. Verklaring gebruik generatieve AI (apart document)")
    doc.add_paragraph("C. Wireframes / mock-ups")
    doc.add_paragraph("D. Bronnenlijst")

    add_prompt(doc,
        "Bronnenlijst: minimaal het XJTU-SY paper (Wang et al., 2018), NLR-website, "
        "AVG en EU AI Act, gebruikte tools/libraries, eventuele wetenschappelijke "
        "literatuur over bearing prognostics.")

    out = "/Users/maxgerbens/Documents/GitHub/nlr-bearing-dashboard/docs/Ontwerpdocument_NLR_Bearing_Dashboard_v0.2.docx"
    doc.save(out)
    print(f"Opgeslagen: {out}")


if __name__ == "__main__":
    build()
