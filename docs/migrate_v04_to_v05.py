"""
Migreert v0.4 naar v0.5.

Wijziging: paragraaf 4.4 (Architectuur) wordt ingevuld met een beschrijving
van de datastroom in zes stappen, elk met: wat er gebeurt, welk artefact
wordt opgeslagen en welke tools worden ingezet.

De rest van het document blijft ongewijzigd.
Versielabel op de voorpagina wordt bijgewerkt naar 0.5.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os, pathlib
_HERE = pathlib.Path(__file__).parent
SRC = str(_HERE / "Ontwerpdocument_NLR_Bearing_Dashboard_v0.4.docx")
DST = str(_HERE / "Ontwerpdocument_NLR_Bearing_Dashboard_v0.5.docx")


def add_para_before(doc, anchor_el, runs):
    """Voeg een alinea in vóór anchor_el. runs = list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold
    anchor_el.addprevious(p._element)
    return p


def add_table_before(doc, anchor_el, headers, rows):
    """Voeg een tabel in vóór anchor_el."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    anchor_el.addprevious(table._element)
    spacer = doc.add_paragraph()
    anchor_el.addprevious(spacer._element)


def main():
    doc = Document(SRC)

    # --- 1. Versie bijwerken op voorpagina -----------------------------------
    for p in doc.paragraphs:
        for run in p.runs:
            if "Versie: 0.4" in run.text:
                run.text = run.text.replace("Versie: 0.4", "Versie: 0.5")

    # --- 2. Invul-prompt van 4.4 lokaliseren ---------------------------------
    prompt_idx = None
    next_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "4.4 Architectuur":
            # De prompt staat direct na de heading
            prompt_idx = i + 1
        if prompt_idx is not None and i > prompt_idx and (
            t.startswith("4.5") or t.startswith("4.6") or t.startswith("5.")
        ):
            next_heading_idx = i
            break

    if prompt_idx is None:
        raise SystemExit("Kon 4.4-prompt niet vinden — script gestopt.")

    # Verwijder de invul-prompt (en eventuele lege regels ertussen)
    paragraphs_to_remove = doc.paragraphs[prompt_idx:next_heading_idx]
    anchor_el = doc.paragraphs[next_heading_idx]._element
    for p in list(paragraphs_to_remove):
        p._element.getparent().remove(p._element)

    # --- 3. Introductie-alinea -----------------------------------------------
    add_para_before(doc, anchor_el, [(
        "De verwerkingspijplijn bestaat uit zes stappen. Elke stap is "
        "onafhankelijk uitvoerbaar: stap 1-3 worden eenmalig gedraaid bij "
        "de eerste setup; stap 4-6 worden bij elke dashboardsessie opnieuw "
        "doorlopen. De tabel hieronder beschrijft per stap wat er gebeurt, "
        "welk artefact wordt opgeslagen en welke tools worden ingezet.", False
    )])

    # --- 4. Datastroom-tabel -------------------------------------------------
    headers = ["Stap", "Wat er gebeurt", "Artefact", "Tools"]
    rows = [
        (
            "1 – Ruwe data\n(data/)",
            "De XJTU-SY dataset staat op schijf als losse CSV-bestanden, "
            "georganiseerd per operating condition en per lager "
            "(data/{conditie}/{lager}/1.csv … N.csv). "
            "Elk bestand bevat één seconde aan vibratiedata: 25.600 samples "
            "op twee kanalen (Horizontal_vibration_signals en "
            "Vertical_vibration_signals). "
            "Er wordt in deze stap niets berekend of gewijzigd; "
            "de bestanden dienen als onveranderlijke bron.",
            "Ongewijzigde CSV-bestanden in data/{conditie}/{lager}/",
            "Geen — pure opslag op schijf.",
        ),
        (
            "2 – Preprocessing",
            "Een Python-script laadt alle CSV's van één lager in één keer in. "
            "Per bestand wordt het bestandsnummer omgezet naar een tijdstempel "
            "(bestandsnummer × 1 s). "
            "Kolomnamen worden gevalideerd; ontbrekende of corrupte bestanden "
            "worden gelogd en overgeslagen. "
            "Het resultaat is één tijdgesorteerde DataFrame per lager.",
            "In-memory pandas DataFrame; geen tussenartefact op schijf "
            "(preprocessing is snel genoeg om bij elke run te herhalen).",
            "Python 3, pandas.",
        ),
        (
            "3 – Feature-extractie\n(features/)",
            "Per meetmoment (= per CSV) worden statistisch-temporele kenmerken "
            "berekend voor zowel het horizontale als het verticale kanaal: "
            "RMS (root mean square, maat voor energieniveau), "
            "kurtosis (maat voor impulsiviteit, verhoogt bij lokale lagerschade), "
            "en crest factor (verhouding piekwaarde / RMS, gevoelig voor "
            "vroege degradatie). "
            "De kenmerken worden per lager opgeslagen als één CSV-bestand; "
            "alle lagers worden vervolgens samengevoegd tot één geconsolideerd "
            "Parquet-bestand.",
            "features/{conditie}/{lager}_features.csv (per lager); "
            "features/features_all.parquet (geconsolideerd, inclusief metadata "
            "zoals conditie, lagernummer, faulttype en faalmoment).",
            "Python 3, NumPy (RMS, piekwaarde), SciPy (kurtosis), pandas.",
        ),
        (
            "4 – Aggregatie",
            "Het geconsolideerde Parquet-bestand wordt aangevuld met "
            "metadata die niet in de ruwe data zit: faulttype per lager "
            "(outer race, inner race, cage, ball of combinatie) en het "
            "bestandsnummer waarop falen is opgetreden. "
            "Deze metadata is afkomstig uit de XJTU-SY dataset-documentatie "
            "en wordt eenmalig als JSON-configuratiebestand bijgehouden. "
            "De samengevoegde tabel vormt de enige databron voor de "
            "visualisatielaag.",
            "features/features_all.parquet (bijgewerkt met metadata-kolommen: "
            "conditie, lager_id, faulttype, fail_idx).",
            "Python 3, pandas, JSON (metadata-config).",
        ),
        (
            "5 – Visualisatielaag",
            "Bij elke gebruikersinteractie (selectie van conditie, lager of "
            "meetmoment) filtert de visualisatielaag de relevante rijen uit "
            "features_all.parquet en berekent de benodigde grafiekinput. "
            "Voor de ruwe signaalgrafiek wordt het bijbehorende CSV-bestand "
            "uit data/ opgehaald. "
            "Er wordt niets naar schijf geschreven; alle bewerkingen zijn "
            "in-memory en duren minder dan 2 seconden (NFE-1).",
            "Geen persistent artefact; in-memory pandas DataFrame en "
            "Plotly Figure-objecten binnen de dashboardsessie.",
            "pandas, Plotly.",
        ),
        (
            "6 – Dashboard-UI",
            "Streamlit host de interactieve gebruikersinterface. "
            "De zijbalk bevat filters voor operating condition (FE-7) en "
            "lagerselectie (FE-1). "
            "Het hoofdpaneel toont: (a) een trendgrafiek van RMS, kurtosis en "
            "crest factor over de levensduur met het falenmoment gemarkeerd "
            "(FE-2, FE-3), (b) een vergelijkingsgrafiek voor meerdere lagers "
            "onder dezelfde conditie (FE-5), en (c) op klik op een datapunt "
            "het bijbehorende ruwe vibratiesignaal (FE-4). "
            "Elke grafiek bevat een korte toelichting (NFE-3) en alle "
            "bediening verloopt via duidelijk gelabelde elementen (NFE-2). "
            "De app draait lokaal via `streamlit run app.py`; "
            "geen externe serverinfrastructuur vereist.",
            "Geen — stateless web-UI; toestand wordt door Streamlit "
            "in-session bijgehouden (st.session_state).",
            "Streamlit, Plotly.",
        ),
    ]
    add_table_before(doc, anchor_el, headers, rows)

    # --- 5. Slotopmerking over uitbreidbaarheid ------------------------------
    add_para_before(doc, anchor_el, [(
        "Uitbreidbaarheid. ",
        True
    ),(
        "Stap 3 kan worden uitgebreid met aanvullende kenmerken "
        "(bv. FFT-piekfrequenties, envelope-analyse) door extra kolommen "
        "toe te voegen aan de feature-CSV's; de downstream stappen hoeven "
        "dan niet aangepast te worden. "
        "Als de dataset in een latere versie groeit, is overstap van Parquet "
        "naar een lichtgewicht database (bv. DuckDB) mogelijk zonder "
        "aanpassing van de visualisatielaag.",
        False
    )])

    doc.save(DST)
    print(f"Opgeslagen: {DST}")


if __name__ == "__main__":
    main()
